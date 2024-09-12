import re
import sys
import yaml
from docx import Document
from docx.shared import Cm
from num2words import num2words
import babel.dates
import datetime


def num2str(num):
    if isinstance(num, str):
        num = int(num)
    return num2words(num, lang='ru')


def load_yaml(yaml_file):
    try:
        with open(yaml_file, 'r', encoding='utf-8') as file:
            data = yaml.load(file, Loader=yaml.FullLoader)
    except yaml.YAMLError as exc:
        print(f"Ошибка при чтении YAML файла: {exc}")
        return None
    return data


class FilesProcessor:
    def __init__(self, yaml_file):
        self.data_dict = load_yaml(yaml_file)
        if self.data_dict is None:
            sys.exit(-1)
        self.instruments = {
            **{key: self.data_dict[key] for key in self.data_dict.keys()},
            **{"num2str": num2str, "babel": babel, "datetime": datetime}
        }
        files_to_processing = self.data_dict["processing"]
        for file_to_processing_key in files_to_processing.keys():
            file_to_processing_data = files_to_processing[file_to_processing_key]
            file_template = file_to_processing_data["template"]
            file_dest = str(self.process_text(file_to_processing_data['dest']))
            if file_dest.endswith(".docx"):
                self.process_word_file(word_file=file_template, output_file=file_dest)

    def process_word_file(self, word_file, output_file):
        try:
            document = Document(word_file)
        except Exception as e:
            print(f"Не удалось открыть файл Word. Ошибка: {str(e)}")
            sys.exit(-1)
        for paragraph in document.paragraphs:
            paragraph.text = self.process_text(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    if i == 1:
                        cell.width = Cm(30)
                    for table_cell_paragraph in cell.paragraphs:
                        table_cell_paragraph.text = self.process_text(table_cell_paragraph.text)
        try:
            document.save(output_file)
        except Exception as e:
            print(f"Не удалось сохранить заполненный файл Word. Ошибка: {str(e)}")

    def process_text(self, text):
        def replace_code(match):
            try:
                return str(eval(match.group(1), self.instruments))
            except Exception as e:
                print(f"Не удалось обработать выражение {match.group(1)}, ошибка: {str(e)}")
                return ""

        return re.sub(r'\{(.*?)\}', replace_code, text)


if __name__ == "__main__":
    yaml_file = 'config.yaml'
    FilesProcessor(yaml_file)
